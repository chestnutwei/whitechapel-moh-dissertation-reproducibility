"""Build Appendix A from the frozen controlled dictionary and final corpus counts."""

from __future__ import annotations

import argparse
import csv
from pathlib import Path
import sys

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

SCRIPT_PATH = Path(__file__).resolve()
PROJECT_ROOT = SCRIPT_PATH.parents[1]
DEFAULT_SUMMARY = PROJECT_ROOT / "outputs" / "tables" / "moh_narrative_keyword_summary.csv"
DEFAULT_REFERENCE = PROJECT_ROOT / "outputs" / "appendices" / "Appendix_A_Controlled_Keyword_Dictionary.docx"
DEFAULT_OUTPUT = PROJECT_ROOT / "outputs" / "appendices" / "Appendix_A_Controlled_Keyword_Dictionary.docx"
DEFAULT_SENSITIVITY = PROJECT_ROOT / "outputs" / "tables" / "technical" / "inequality_marker_sensitivity_summary.csv"

sys.path.insert(0, str(SCRIPT_PATH.parent))
from moh_keyword_dictionary import (  # noqa: E402
    TERM_GROUPS,
    LEXICAL_VARIANTS,
    RETRIEVAL_TYPE,
    RATIONALE,
    OCR_EXCLUSIONS,
)

DISPLAY_CATEGORY = {"disease": "Disease", "sanitary": "Sanitation", "housing": "Housing", "governance": "Governance", "inequality": "Inequality"}
CATEGORY_ORDER = ["disease", "sanitary", "housing", "governance", "inequality"]


def set_run_font(run, size: float, *, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Times New Roman")
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Times New Roman")
    normal._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(3)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.0

    for style_name, size in (("Heading 1", 16), ("Heading 2", 13)):
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Times New Roman")
        style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(4)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.keep_with_next = True

    if "Appendix Caption" not in [style.name for style in doc.styles]:
        caption = doc.styles.add_style("Appendix Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        caption = doc.styles["Appendix Caption"]
    caption.font.name = "Times New Roman"
    caption.font.size = Pt(10)
    caption.font.italic = True
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(4)


def clear_body(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def load_counts(path: Path) -> dict[tuple[str, str], tuple[int, float, int]]:
    counts = {}
    with path.open(newline="", encoding="utf-8") as handle:
        for row in csv.DictReader(handle):
            counts[(row["category"], row["term"])] = (
                int(row["raw_count"]),
                float(row["frequency_per_10000_words"]),
                int(row["narrative_word_count"]),
            )
    return counts


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def format_cell(cell, *, bold: bool = False, center: bool = False) -> None:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        for run in paragraph.runs:
            set_run_font(run, 8.2, bold=bold)


def set_cell_width(cell, inches: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def build_appendix(reference: Path, summary: Path, output: Path, sensitivity: Path = DEFAULT_SENSITIVITY) -> None:
    counts = load_counts(summary)
    sensitivity_rows = list(csv.DictReader(sensitivity.open(newline="", encoding="utf-8"))) if sensitivity.exists() else []
    rows = []
    for category in CATEGORY_ORDER:
        for canonical in TERM_GROUPS[category]:
            raw, freq, words = counts[(category, canonical)]
            variants = "; ".join(LEXICAL_VARIANTS.get(canonical, [])) or "—"
            retrieval_type = RETRIEVAL_TYPE.get(canonical, "direct")
            rationale = RATIONALE.get(canonical, "Direct retrieval term within the stated analytical category.")
            rows.append((DISPLAY_CATEGORY[category], canonical, variants, retrieval_type, rationale, raw, freq, words))

    if len(rows) != 55:
        raise ValueError(f"Expected 55 canonical dictionary entries; found {len(rows)}")
    if {r[7] for r in rows} != {51428}:
        raise ValueError("Unexpected narrative word count in summary table")

    doc = Document(reference) if reference.exists() else Document()
    clear_body(doc)
    configure_styles(doc)
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    title = doc.add_paragraph(style="Heading 1")
    set_run_font(title.add_run("Appendix A. Controlled Keyword Dictionary"), 16, bold=True)

    heading = doc.add_paragraph(style="Heading 2")
    set_run_font(heading.add_run("A.1 Scope and matching rules"), 13, bold=True)

    p = doc.add_paragraph()
    set_run_font(p.add_run(
        "This appendix records the frozen controlled dictionary used for retrieval and counting in the final 275-passage, 51,428-word narrative corpus. "
        "The dictionary contains 55 canonical analytical entries across five researcher-defined categories: disease (13), sanitation (14), housing (9), governance (10), and inequality (9). "
        "A canonical entry may correspond to more than one source-text surface form."
    ), 10.5)

    p = doc.add_paragraph()
    set_run_font(p.add_run(
        "A limited lexical-normalisation layer maps only a small, frozen set of source-verified lexical variants observed in the reports to their canonical entries. "
        "These variants cover hyphenation, spacing and number differences. The layer does not perform fuzzy matching or open-ended OCR correction. "
        "OCR-induced misreadings are documented during source verification rather than added to the alias rules."
    ), 10.5)

    p = doc.add_paragraph()
    set_run_font(p.add_run(
        "The inequality category distinguishes direct socioeconomic vocabulary from indirect administrative or institutional markers. "
        "The terms 'lodgers' and 'infirmary' broaden candidate retrieval but do not automatically establish inequality; their hits require contextual review. "
        "Material conditions such as lodging-house occupation, crowding, limited storage capacity and displacement are interpreted through close reading rather than encoded as automatic inequality synonyms."
    ), 10.5)

    heading = doc.add_paragraph(style="Heading 2")
    set_run_font(heading.add_run("A.2 Dictionary, lexical variants and corpus counts"), 13, bold=True)

    caption = doc.add_paragraph(style="Appendix Caption")
    set_run_font(caption.add_run("Table A1. Frozen controlled dictionary, source-verified lexical variants and final-corpus hit counts"), 10, italic=True)

    headers = ["Category", "Canonical entry", "Source-verified lexical variants", "Retrieval type", "Rationale", "Raw hits", "Hits per 10,000 words"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    widths = [0.8, 1.25, 2.15, 1.35, 3.35, 0.65, 1.05]
    for i, text in enumerate(headers):
        table.rows[0].cells[i].text = text
        shade_cell(table.rows[0].cells[i], "E7E6E6")
        format_cell(table.rows[0].cells[i], bold=True, center=True)
        set_cell_width(table.rows[0].cells[i], widths[i])
    repeat_header(table.rows[0])

    for category, canonical, variants, retrieval_type, rationale, raw, freq, _words in rows:
        row = table.add_row()
        values = [category, canonical, variants, retrieval_type, rationale, str(raw), f"{freq:.2f}"]
        for i, value in enumerate(values):
            row.cells[i].text = value
            format_cell(row.cells[i], center=i in (5, 6))
            set_cell_width(row.cells[i], widths[i])
        prevent_row_split(row)

    heading = doc.add_paragraph(style="Heading 2")
    set_run_font(heading.add_run("A.3 Inequality-marker sensitivity check"), 13, bold=True)

    p = doc.add_paragraph()
    set_run_font(p.add_run(
        "Because 'lodgers' and 'infirmary' are context-dependent indirect markers, a sensitivity check compares the full inequality retrieval layer with direct socioeconomic vocabulary only. "
        "The full layer is retained for candidate retrieval, while the direct-only series is used for the descriptive inequality column in Figure 3 so that context-dependent institutional language does not inflate the year-by-year frequency pattern."
    ), 10.5)

    if sensitivity_rows:
        headers2 = ["Scope", "Candidate records", "Distinct passages", "Raw hits"]
        table2 = doc.add_table(rows=1, cols=len(headers2))
        table2.autofit = False
        widths2 = [2.0, 1.5, 1.5, 1.2]
        for i, text in enumerate(headers2):
            table2.rows[0].cells[i].text = text
            shade_cell(table2.rows[0].cells[i], "E7E6E6")
            format_cell(table2.rows[0].cells[i], bold=True, center=True)
            set_cell_width(table2.rows[0].cells[i], widths2[i])
        repeat_header(table2.rows[0])
        for sr in sensitivity_rows:
            row = table2.add_row()
            scope_label = "Direct + indirect markers" if sr["scope"] == "all" else "Direct socioeconomic entries only"
            vals = [scope_label, sr["candidate_records"], sr["distinct_passages"], sr["raw_hits"]]
            for i, value in enumerate(vals):
                row.cells[i].text = str(value)
                format_cell(row.cells[i], center=i > 0)
                set_cell_width(row.cells[i], widths2[i])
            prevent_row_split(row)

    p = doc.add_paragraph()
    set_run_font(p.add_run(
        "The difference is material: the indirect markers add 21 raw hits and 12 distinct candidate passages. This confirms that they are useful for retrieval breadth but should not be treated as validated inequality prevalence without contextual review."
    ), 10.5)

    heading = doc.add_paragraph(style="Heading 2")
    set_run_font(heading.add_run("A.4 OCR exclusion example"), 13, bold=True)
    for form, reason in OCR_EXCLUSIONS.items():
        p = doc.add_paragraph()
        set_run_font(p.add_run(f"{form}: "), 10.5, bold=True)
        set_run_font(p.add_run(reason), 10.5)

    note = doc.add_paragraph()
    set_run_font(note.add_run("Note: "), 9.5, bold=True)
    set_run_font(note.add_run(
        "Raw counts are retrieval diagnostics rather than validated measures of thematic prevalence. In the project CSV files, the dissertation category 'sanitation' is stored under the machine-readable label 'sanitary'."
    ), 9.5)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    print(f"Wrote {output}")
    print(f"Canonical dictionary entries: {len(rows)}")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("--reference", type=Path, default=DEFAULT_REFERENCE)
    parser.add_argument("--summary", type=Path, default=DEFAULT_SUMMARY)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    parser.add_argument("--sensitivity", type=Path, default=DEFAULT_SENSITIVITY)
    return parser.parse_args()


if __name__ == "__main__":
    args = parse_args()
    build_appendix(args.reference, args.summary, args.output, args.sensitivity)
